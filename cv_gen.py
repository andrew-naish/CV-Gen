import yaml
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import os
import argparse
from datetime import datetime

## Init

# Parse command-line arguments for YAML path
parser = argparse.ArgumentParser(description='Generate CV from YAML')
parser.add_argument('--yaml', type=str, default='default.yaml', help='Path to YAML file (default: default.yaml)')
args = parser.parse_args()
yaml_path = args.yaml
yaml_path_filename = os.path.basename(yaml_path)

def add_heading(doc, heading, level=1, space_before=Pt(0)):
    heading = doc.add_heading(heading, level)
    heading.paragraph_format.space_before = space_before
    return heading

def add_heading_list(doc, heading_text, items):
    add_heading(doc, heading_text)
    paragraph = doc.add_paragraph()
    item_count = len(items)
    for idx, (key, value) in enumerate(items.items()):
        run_key = paragraph.add_run(key)
        run_key.bold = True
        paragraph.add_run(": ")
        paragraph.add_run(str(value))
        if idx < item_count - 1:
            paragraph.add_run("\n")

def add_heading_text(doc, heading_text, text):
    add_heading(doc, heading_text)
    paragraph = doc.add_paragraph()
    paragraph.add_run(text)

def add_heading_bulleted_category_list(doc, heading_text, items):
    """
    Adds a heading and a list of categories, where each key is bold and each value is a comma-separated list on the same line. The values are sorted alphabetically.
    """
    add_heading(doc, heading_text)
    for key, value in items.items():
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(str(key))
        run.bold = True
        p.add_run(": ")
        if isinstance(value, list):
            sorted_values = sorted(str(v) for v in value)
            p.add_run(", ".join(sorted_values))
        else:
            p.add_run(str(value))

## Main

# Load YAML data
if not os.path.isfile(yaml_path):
    print(f"Error: YAML file '{yaml_path}' does not exist.")
    exit(1)

with open(yaml_path, 'r', encoding='utf-8') as f:
    cv_data = yaml.safe_load(f)

# create doc and set options
doc = Document()

# Set default body font to Aptos
style = doc.styles['Normal']
font = style.font
font.name = 'Aptos'
font.size = Pt(11)
font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')

# Set default headings font to Aptos Display
for heading_level in range(1, 10):
    style_name = f'Heading {heading_level}'
    if style_name in doc.styles:
        heading_style = doc.styles[style_name]
        heading_font = heading_style.font
        heading_font.name = 'Aptos Display'
        heading_font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos Display')

# Update core properties
doc.core_properties.title = f"CV {cv_data['name']}"
doc.core_properties.author = cv_data['name']
doc.core_properties.comments = ""

# Title
doc.add_heading(cv_data['name'], 0)

# Contact Info
if 'contact_information' in cv_data:
    add_heading_list(doc, "Contact Information", cv_data['contact_information'])

# Links
if 'links' in cv_data:
    add_heading_list(doc, "Links", cv_data['links'])

# Personal Summary
if 'personal_summary' in cv_data:
    add_heading_text(doc, "Personal Summary", cv_data['personal_summary'])

# Key Skills
if 'key_skills' in cv_data:
    add_heading_bulleted_category_list(doc, "Key Skills", cv_data['key_skills'])

# Professional Experience
## TODO: Add professional experience section

# Certifications
## TODO: Add certifications section

# Education
## TODO: Add education section

## Output

# generate known name parts
output_name_prefix = "CV"
output_name_suffix = datetime.now().strftime("%Y-%m")

# generate dynamic name parts
if yaml_path_filename == "default.yaml":
    output_name_mid = cv_data['name']
else:
    output_name_mid = os.path.splitext(yaml_path_filename)[0]

# smoosh it all together
output_name = f"{output_name_prefix} - {output_name_mid} - {output_name_suffix}"

# ensure the output directory exists
output_dir = "output"
os.makedirs(output_dir, exist_ok=True)

# save the document
doc.save(os.path.join(output_dir, f"{output_name}.docx"))
